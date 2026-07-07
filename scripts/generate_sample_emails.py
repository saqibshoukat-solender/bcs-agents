#!/usr/bin/env python3
"""
Generate 15 sample customer emails using Casey's exact email logic.
Read-only — no emails sent, no database writes, no agent runs triggered.

Job selection targets all 4 email scenarios so the reviewer sees a
representative mix (in_progress, not_started, invoice_reminder, new_job_intro).

Run inside the casey container:
    docker compose run --rm casey python scripts/generate_sample_emails.py

Output files are written inside the container at /opt/bcs-agents/.
To retrieve them after the run (drop --rm, then docker cp):
    docker compose run --name casey_sample casey python scripts/generate_sample_emails.py
    docker cp casey_sample:/opt/bcs-agents/sample_emails.txt ./sample_emails.txt
    docker cp casey_sample:/opt/bcs-agents/sample_emails.docx ./sample_emails.docx
    docker rm casey_sample

Or mount a host directory and write directly to disk:
    docker compose run --rm -v /opt/bcs-agents:/opt/bcs-agents casey \
        python scripts/generate_sample_emails.py
"""
import os
import re
import sys
from datetime import datetime, date, timezone

# Project root on path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from dotenv import load_dotenv
load_dotenv()

from db.state_store import get_all_active_jobs, get_pm_list, get_email_history
from integrations.quickbooks import get_invoice_status_for_customer
from integrations.sheets import parse_latest_date
from agents.casey.email_composer import compose_customer_update_email
from agents.casey.main import normalize_pm_name
from utils.logger import get_logger

logger = get_logger("generate_sample_emails")

OUTPUT_TXT  = os.getenv("SAMPLE_EMAILS_TXT",  "/opt/bcs-agents/sample_emails.txt")
OUTPUT_DOCX = os.getenv("SAMPLE_EMAILS_DOCX", "/opt/bcs-agents/sample_emails.docx")
TOTAL = 15

# How many jobs to target per scenario (must sum to TOTAL)
_TARGETS = {
    "in_progress":      4,
    "not_started":      4,
    "invoice_reminder": 3,
    "new_job_intro":    4,
}


# ---------------------------------------------------------------------------
# Helpers (mirrors of casey/main.py — inlined so no agent startup side-effects)
# ---------------------------------------------------------------------------

def _determine_scenario(job: dict) -> str:
    if job.get("sheet_tab") == "to_start":
        return "not_started"
    qb_status = job.get("qb_invoice_status")
    if qb_status and qb_status.get("days_overdue", 0) >= 60:
        return "invoice_reminder"
    if job.get("assigned_crew_sub", "").strip():
        return "in_progress"
    start_str = job.get("start_date", "").strip()
    if start_str:
        try:
            start = datetime.strptime(start_str, "%Y-%m-%d").date()
            if start <= date.today():
                return "in_progress"
        except Exception:
            pass
    return "not_started"


def _get_pm_email(pm_name: str, pm_list: list) -> str:
    normalized = normalize_pm_name(pm_name)
    for pm in pm_list:
        if normalize_pm_name(pm.get("full_name", "")) == normalized:
            return pm.get("email", "")
    return ""


def _pm_known(pm_name: str, pm_list: list) -> bool:
    """True if pm_name resolves to a real entry in pm_config."""
    if not pm_name.strip():
        return False
    normalized = normalize_pm_name(pm_name)
    return any(normalize_pm_name(pm.get("full_name", "")) == normalized for pm in pm_list)


def _days_since(d) -> "int | None":
    if not d:
        return None
    if isinstance(d, str):
        try:
            d = datetime.strptime(d, "%Y-%m-%d").date()
        except Exception:
            return None
    try:
        return (date.today() - d).days
    except Exception:
        return None


def _strip_html(html: str) -> str:
    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.IGNORECASE)
    text = re.sub(r"<p[^>]*>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</p>", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Job eligibility and scenario pre-classification
# ---------------------------------------------------------------------------

def _is_eligible(job: dict, pm_list: list) -> bool:
    """
    A job must have a real customer email, a valid PM in pm_config,
    and must not be addressed to an internal BCS address.
    """
    email = job.get("customer_email", "").strip()
    if not email:
        return False
    if "bluecollarscholars.net" in email.lower():
        return False
    pm_name = job.get("pm_name", "").strip()
    if not _pm_known(pm_name, pm_list):
        return False
    return True


def _pre_classify(job: dict) -> str:
    """
    Classify a job's likely scenario using DB fields only (no QB call).
    invoice_reminder is simulated: to_collect non-empty + deposit > 60 days old.
    new_job_intro check mirrors the override in agents/casey/main.py.
    """
    # new_job_intro: never emailed, fresh deposit (within 14 days)
    if (job.get("last_customer_update_sent") is None
            and job.get("next_scheduled_update") is None):
        deposit = parse_latest_date(job.get("deposit_date", ""))
        if deposit and 0 <= (date.today() - deposit).days <= 14:
            return "new_job_intro"

    # invoice_reminder simulation (QB may not be available)
    to_collect = job.get("to_collect", "").strip()
    if to_collect:
        deposit = parse_latest_date(job.get("deposit_date", ""))
        if deposit and (date.today() - deposit).days >= 60:
            return "invoice_reminder"

    # in_progress
    if job.get("assigned_crew_sub", "").strip():
        return "in_progress"
    start_str = job.get("start_date", "").strip()
    if start_str:
        try:
            start = datetime.strptime(start_str, "%Y-%m-%d").date()
            if start <= date.today():
                return "in_progress"
        except Exception:
            pass

    return "not_started"


# ---------------------------------------------------------------------------
# Job selection
# ---------------------------------------------------------------------------

def select_jobs(pm_list: list, total: int = TOTAL) -> list[dict]:
    """
    Load all active jobs, filter for eligibility, and pick a balanced set
    that covers all 4 email scenarios. Falls back to any surplus eligible
    jobs if a bucket has fewer than its target.
    """
    all_jobs = get_all_active_jobs()

    # Sort by next_scheduled_update ASC NULLS LAST (Casey's priority queue)
    all_jobs.sort(key=lambda j: (
        j.get("next_scheduled_update") is None,
        j.get("next_scheduled_update") or date.max,
    ))

    eligible = [j for j in all_jobs if _is_eligible(j, pm_list)]

    # Tag and bucket
    buckets: dict[str, list[dict]] = {
        "in_progress": [], "not_started": [],
        "invoice_reminder": [], "new_job_intro": [],
    }
    for job in eligible:
        sc = _pre_classify(job)
        job["_pre_scenario"] = sc
        buckets[sc].append(job)

    print(f"  Eligible jobs by pre-classified scenario:")
    for sc, jobs in buckets.items():
        target = _TARGETS[sc]
        print(f"    {sc:<20} {len(jobs):>3} available  (target: {target})")

    # Pick up to target from each bucket (preserving order = priority)
    selected: list[dict] = []
    selected_ids: set = set()

    for sc, target in _TARGETS.items():
        taken = 0
        for job in buckets[sc]:
            if taken >= target:
                break
            jid = job.get("id") or job["client_name"]
            if jid not in selected_ids:
                selected.append(job)
                selected_ids.add(jid)
                taken += 1

    # Fill any remaining slots from all eligible jobs not already picked
    if len(selected) < total:
        for job in eligible:
            if len(selected) >= total:
                break
            jid = job.get("id") or job["client_name"]
            if jid not in selected_ids:
                selected.append(job)
                selected_ids.add(jid)

    return selected[:total]


# ---------------------------------------------------------------------------
# Per-job processing
# ---------------------------------------------------------------------------

def build_email_record(job: dict, idx: int, total: int, pm_list: list) -> dict:
    client_name  = job["client_name"]
    pm_name      = job.get("pm_name", "")
    pre_scenario = job.get("_pre_scenario", "")

    record = {
        "idx":               idx,
        "total":             total,
        "client_name":       client_name,
        "pm_name":           pm_name,
        "job_type":          job.get("job_type", ""),
        "next_update":       job.get("next_scheduled_update"),
        "customer_email":    job.get("customer_email", ""),
        "pm_email":          "",
        "subject":           "",
        "body_plain":        "",
        "scenario":          "",
        "escalation":        False,
        "escalation_reason": "",
        "comms_hold":        False,
        "comms_hold_reason": "",
        "notes":             [],
    }

    # ── Comms hold — include in output but note it ───────────────────────────
    if job.get("comms_hold"):
        reason = job.get("comms_hold_reason") or "no reason recorded"
        record.update({
            "comms_hold":        True,
            "comms_hold_reason": reason,
            "scenario":          "skipped",
            "subject":           "(SKIPPED — comms hold active)",
            "body_plain":        "(This job is on comms hold — Casey would not send an email.)",
        })
        record["notes"].append(f"COMMS HOLD: {reason}")
        return record

    # ── QB invoice status ────────────────────────────────────────────────────
    qb_status = None
    try:
        qb_status = get_invoice_status_for_customer(client_name)
        if not qb_status or not qb_status.get("found"):
            record["notes"].append("QB invoice not found for this customer")
    except Exception as e:
        record["notes"].append(f"QB lookup failed: {e}")
    job["qb_invoice_status"] = qb_status

    # ── Scenario determination ───────────────────────────────────────────────
    scenario = _determine_scenario(job)

    # new_job_intro: override when QB didn't change the classification
    # (mirrors the check in agents/casey/main.py that runs after _determine_scenario)
    if pre_scenario == "new_job_intro" and scenario not in ("invoice_reminder",):
        scenario = "new_job_intro"
        record["notes"].append(
            "new_job_intro: first Casey email for this job, deposit received within 14 days"
        )

    # invoice_reminder: inject simulated status when QB is not confirming overdue
    if pre_scenario == "invoice_reminder" and scenario != "invoice_reminder":
        scenario = "invoice_reminder"
        record["notes"].append(
            "invoice_reminder simulated — QB not confirming overdue; "
            "selected because to_collect is set and deposit is >60 days old"
        )

    record["scenario"] = scenario

    # ── Escalation check (note only — email still generated) ─────────────────
    history = get_email_history(client_name, pm_name)
    last_contact = history.last_sent_at if history and history.last_sent_at else None
    days_since = _days_since(last_contact)
    if days_since is not None and days_since >= 14:
        record["escalation"]        = True
        record["escalation_reason"] = f"No PM contact in {days_since} days"
        record["notes"].append(f"WOULD ESCALATE: {record['escalation_reason']}")

    # ── PM email ─────────────────────────────────────────────────────────────
    pm_email = _get_pm_email(pm_name, pm_list)
    record["pm_email"] = pm_email
    if not pm_email:
        record["notes"].append(f"PM email not found in pm_config for '{pm_name}'")

    # ── Compose — same LLM call Casey makes, same fallback template ──────────
    email_snippets = (history.email_snippets or "") if history else ""

    composed = compose_customer_update_email(
        customer_name=client_name,
        pm_name=pm_name,
        job_type=job.get("job_type", ""),
        scenario=scenario,
        contractor=job.get("assigned_crew_sub", ""),
        notes=job.get("pm_communication_history", "") or job.get("pm_notes", ""),
        to_collect=job.get("to_collect", ""),
        job_description=job.get("job_description", ""),
        complaint_note=job.get("complaint_note", ""),
        client_mood=job.get("client_mood", ""),
        total_project=job.get("total_project", ""),
        estimator_name=job.get("estimator_name", ""),
        sheet_tab=job.get("sheet_tab", ""),
        email_history=email_snippets,
    )

    record["subject"]    = composed.get("subject", "")
    record["body_plain"] = _strip_html(composed.get("body_html", ""))
    return record


# ---------------------------------------------------------------------------
# Text file output
# ---------------------------------------------------------------------------

def write_txt(records: list[dict], path: str) -> None:
    scenario_counts: dict[str, int] = {}
    escalation_count = 0
    lines: list[str] = []

    for rec in records:
        W = 80
        lines += [
            "=" * W,
            f"EMAIL #{rec['idx']} of {rec['total']}",
            "=" * W,
            f"Customer:     {rec['client_name']}",
            f"PM:           {rec['pm_name'] or '(none)'}",
            f"Job Type:     {rec['job_type'] or '(none)'}",
            f"Scenario:     {rec['scenario']}",
            f"Next Update:  {rec['next_update'].isoformat() if rec['next_update'] else '(not set)'}",
            f"Escalation:   {'YES — ' + rec['escalation_reason'] if rec['escalation'] else 'NO'}",
            f"Note:         {'; '.join(rec['notes']) if rec['notes'] else ''}",
            "",
            f"FROM:    {rec['pm_email'] or '(PM email unknown)'}",
            f"TO:      {rec['customer_email']}",
            f"SUBJECT: {rec['subject']}",
            "",
            "BODY:",
            rec["body_plain"],
            "",
            "-" * W,
            "REVIEWER NOTES: " + "_" * 63,
            "_" * W,
            "=" * W,
            "",
            "",
        ]

        sc = rec["scenario"]
        scenario_counts[sc] = scenario_counts.get(sc, 0) + 1
        if rec["escalation"]:
            escalation_count += 1

    lines += [
        "=" * 80,
        "SUMMARY",
        "=" * 80,
        f"Total emails generated: {len(records)}",
        "Scenarios breakdown:",
    ]
    for sc in ("in_progress", "not_started", "invoice_reminder", "new_job_intro", "skipped"):
        lines.append(f"  - {sc}: {scenario_counts.get(sc, 0)}")
    lines += [
        f"Escalation flags: {escalation_count} jobs would have been escalated",
        f"Generated at: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}",
        "",
    ]

    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"  Text  → {path}")


# ---------------------------------------------------------------------------
# Word document output
# ---------------------------------------------------------------------------

def write_docx(records: list[dict], path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    def _meta(label: str, value: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_l = p.add_run(f"{label:<14}")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_v = p.add_run(value)
        run_v.font.size = Pt(10)

    def _divider():
        p = doc.add_paragraph("─" * 72)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        for run in p.runs:
            run.font.size = Pt(8)

    def _reviewer_block():
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("REVIEWER NOTES:")
        run.bold = True
        run.font.size = Pt(10)
        for _ in range(3):
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(8)
            run_u = line.add_run("_" * 88)
            run_u.font.size = Pt(10)

    for i, rec in enumerate(records):
        if i > 0:
            doc.add_page_break()

        heading = doc.add_heading(
            f"Email #{rec['idx']} of {rec['total']}  —  {rec['client_name']}",
            level=1,
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        _divider()

        next_upd = rec["next_update"]
        _meta("Customer:",    rec["client_name"])
        _meta("PM:",          rec["pm_name"] or "(none)")
        _meta("Job Type:",    rec["job_type"] or "(none)")
        _meta("Scenario:",    rec["scenario"])
        _meta("Next Update:", next_upd.isoformat() if next_upd else "(not set)")
        _meta("Escalation:",  ("YES — " + rec["escalation_reason"]) if rec["escalation"] else "NO")
        if rec["notes"]:
            _meta("Note:", "; ".join(rec["notes"]))

        _divider()

        _meta("FROM:",    rec["pm_email"] or "(PM email unknown)")
        _meta("TO:",      rec["customer_email"])
        _meta("SUBJECT:", rec["subject"])

        doc.add_paragraph()

        p_label = doc.add_paragraph()
        p_label.paragraph_format.space_after = Pt(4)
        run_lbl = p_label.add_run("BODY:")
        run_lbl.bold = True
        run_lbl.font.size = Pt(10)

        for line in rec["body_plain"].split("\n"):
            p = doc.add_paragraph(line or " ")
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(10)

        _reviewer_block()

    # Summary page
    doc.add_page_break()
    doc.add_heading("Summary", level=1)

    scenario_counts: dict[str, int] = {}
    escalation_count = 0
    for rec in records:
        sc = rec["scenario"]
        scenario_counts[sc] = scenario_counts.get(sc, 0) + 1
        if rec["escalation"]:
            escalation_count += 1

    doc.add_paragraph(f"Total emails generated: {len(records)}")
    doc.add_paragraph("Scenarios breakdown:")
    for sc in ("in_progress", "not_started", "invoice_reminder", "new_job_intro", "skipped"):
        doc.add_paragraph(f"    •  {sc}: {scenario_counts.get(sc, 0)}")
    doc.add_paragraph(f"Escalation flags: {escalation_count} jobs would have been escalated")
    doc.add_paragraph(
        f"Generated at: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}"
    )

    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    print(f"  Word  → {path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    print("Loading PM config and jobs from database...")
    pm_list = get_pm_list()
    jobs = select_jobs(pm_list, TOTAL)

    if not jobs:
        print("ERROR: No eligible jobs found in database.")
        sys.exit(1)

    print(f"\nSelected {len(jobs)} jobs. Composing emails (LLM calls in progress)...\n")
    records: list[dict] = []
    for i, job in enumerate(jobs, 1):
        sc_label = job.get("_pre_scenario", "?")
        print(f"  [{i:>2}/{len(jobs)}] {job['client_name']:<35} scenario={sc_label}  pm={job.get('pm_name','')}")
        records.append(build_email_record(job, i, len(jobs), pm_list))

    print("\nWriting output files...")
    write_txt(records, OUTPUT_TXT)
    write_docx(records, OUTPUT_DOCX)
    print("\nDone.")


if __name__ == "__main__":
    main()
